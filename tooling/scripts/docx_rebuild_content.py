#!/usr/bin/env python3
"""
docx_rebuild_content.py - 保留 DOCX 格式，完全替換正文內容

策略：
1. 保留 DOCX 的封面段落（中文摘要之前）
2. 保留 styles.xml, numbering.xml, header, footer 等格式定義
3. 清空中文摘要之後的所有段落
4. 按照 TARGET_MD 的內容重新建立段落，使用 DOCX 現有樣式

用法：
    python docx_rebuild_content.py template.docx content.md output.docx
"""

import argparse
import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from markdown_it import MarkdownIt


def get_para_text(para) -> str:
    return "".join(node.text or "" for node in para._element.iter(qn("w:t"))).strip()


def get_para_style(para) -> str:
    return para.style.name if para.style else "Normal"


def set_para_text(para, new_text: str) -> bool:
    t_elements = list(para._element.iter(qn("w:t")))
    if not t_elements:
        return False
    for t in t_elements:
        t.text = ""
    t_elements[0].text = new_text
    return True


def clone_para_after(para, new_text: str, style_name: str = None):
    """在段落後插入複製的段落"""
    new_elem = deepcopy(para._element)
    para._element.addnext(new_elem)
    t_elements = list(new_elem.iter(qn("w:t")))
    if t_elements:
        for t in t_elements:
            t.text = ""
        t_elements[0].text = new_text
    return new_elem


def parse_md(md_path: str) -> list[dict]:
    """解析 Markdown，返回段落列表"""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    md = MarkdownIt()
    tokens = md.parse(content)

    result = []
    i = 0
    while i < len(tokens):
        token = tokens[i]

        if token.type == "heading_open":
            level = int(token.tag[1])  # h1 -> 1, h2 -> 2
            # 下一個 token 是 inline
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                text = tokens[i + 1].content
                result.append({
                    "type": "heading",
                    "level": level,
                    "text": text
                })
            i += 3  # heading_open, inline, heading_close
        elif token.type == "paragraph_open":
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                text = tokens[i + 1].content
                # 跳過圖片標記
                if text.startswith("!["):
                    i += 3
                    continue
                result.append({
                    "type": "paragraph",
                    "text": text
                })
            i += 3
        elif token.type == "bullet_list_open" or token.type == "ordered_list_open":
            # 處理列表
            list_type = "bullet" if token.type == "bullet_list_open" else "ordered"
            i += 1
            while i < len(tokens) and tokens[i].type not in ("bullet_list_close", "ordered_list_close"):
                if tokens[i].type == "list_item_open":
                    i += 1
                    while i < len(tokens) and tokens[i].type != "list_item_close":
                        if tokens[i].type == "paragraph_open":
                            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                                text = tokens[i + 1].content
                                prefix = "• " if list_type == "bullet" else ""
                                result.append({
                                    "type": "list_item",
                                    "text": prefix + text
                                })
                            i += 3
                        else:
                            i += 1
                else:
                    i += 1
            i += 1
        elif token.type == "table_open":
            # 處理表格 - 簡化處理，將表格作為文字段落
            table_content = []
            i += 1
            while i < len(tokens) and tokens[i].type != "table_close":
                if tokens[i].type == "inline":
                    table_content.append(tokens[i].content)
                i += 1
            if table_content:
                result.append({
                    "type": "table",
                    "text": " | ".join(table_content[:20])  # 簡化
                })
            i += 1
        elif token.type == "hr":
            # 跳過分隔線
            i += 1
        elif token.type == "fence" or token.type == "code_block":
            # 程式碼區塊
            result.append({
                "type": "code",
                "text": token.content[:500] if token.content else ""
            })
            i += 1
        else:
            i += 1

    return result


def find_style_template(doc, style_type: str):
    """找到特定類型的樣式模板段落"""
    for para in doc.paragraphs:
        text = get_para_text(para)
        style = get_para_style(para)

        if style_type == "heading1" and "Heading 1" in style:
            return para
        elif style_type == "heading2" and "Heading 2" in style:
            return para
        elif style_type == "heading3" and "Heading 3" in style:
            return para
        elif style_type == "normal" and text and "Heading" not in style:
            return para
    return None


def main():
    parser = argparse.ArgumentParser(description="保留 DOCX 格式，完全替換正文內容")
    parser.add_argument("docx", type=Path, help="原始 DOCX")
    parser.add_argument("md", type=Path, help="內容來源 MD")
    parser.add_argument("output", type=Path, help="輸出 DOCX")
    parser.add_argument("--start-marker", default="中文摘要", help="正文開始標記")
    parser.add_argument("--dry-run", action="store_true", help="只顯示將執行的操作")
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"錯誤：找不到 {args.docx}", file=sys.stderr)
        return 1

    if not args.md.exists():
        print(f"錯誤：找不到 {args.md}", file=sys.stderr)
        return 1

    # 解析 MD
    print(f"解析 MD：{args.md}")
    md_paras = parse_md(str(args.md))
    print(f"  - 總段落數：{len(md_paras)}")

    # 找到 MD 中的開始標記位置
    md_start_idx = None
    for i, p in enumerate(md_paras):
        if p["text"] == args.start_marker:
            md_start_idx = i
            break

    if md_start_idx is None:
        print(f"警告：MD 中找不到開始標記 '{args.start_marker}'，使用第一段")
        md_start_idx = 0
    else:
        print(f"  - MD 開始位置：{md_start_idx}")

    # 只取開始標記之後的段落（含開始標記）
    md_paras = md_paras[md_start_idx:]
    print(f"  - 正文段落數：{len(md_paras)}")

    if args.dry_run:
        print("\n=== Dry Run ===")
        for i, p in enumerate(md_paras[:30]):
            print(f"  {i}. [{p['type']}] {p['text'][:60]}...")
        if len(md_paras) > 30:
            print(f"  ... 還有 {len(md_paras) - 30} 項")
        return 0

    # 載入 DOCX
    print(f"載入 DOCX：{args.docx}")
    doc = Document(str(args.docx))

    # 找到正文開始位置
    start_idx = None
    for i, para in enumerate(doc.paragraphs):
        if get_para_text(para) == args.start_marker:
            start_idx = i
            break

    if start_idx is None:
        print(f"錯誤：找不到開始標記 '{args.start_marker}'", file=sys.stderr)
        return 1

    print(f"  - 正文開始位置：{start_idx}")

    # 收集樣式模板
    heading1_template = None
    heading2_template = None
    heading3_template = None
    normal_template = None

    for para in doc.paragraphs[start_idx:]:
        style = get_para_style(para)
        text = get_para_text(para)

        if not text:
            continue

        if "Heading 1" in style and not heading1_template:
            heading1_template = para
        elif "Heading 2" in style and not heading2_template:
            heading2_template = para
        elif "Heading 3" in style and not heading3_template:
            heading3_template = para
        elif "Heading" not in style and not normal_template:
            normal_template = para

    print(f"  - 找到樣式模板：H1={heading1_template is not None}, H2={heading2_template is not None}, H3={heading3_template is not None}, Normal={normal_template is not None}")

    # 如果沒有找到所有模板，使用 fallback
    if not normal_template:
        normal_template = doc.paragraphs[start_idx]
    if not heading1_template:
        heading1_template = normal_template
    if not heading2_template:
        heading2_template = normal_template
    if not heading3_template:
        heading3_template = normal_template

    # 刪除正文段落（保留開始標記）
    body = doc._body._body
    paras_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        if i > start_idx:
            paras_to_remove.append(para._element)

    for elem in paras_to_remove:
        body.remove(elem)

    print(f"  - 已刪除 {len(paras_to_remove)} 個段落")

    # 插入新內容 - 使用 body.append() 在末尾追加
    insert_count = 0

    # 跳過第一個元素（中文摘要標題已經在 DOCX 中）
    for md_para in md_paras[1:]:
        text = md_para["text"]
        para_type = md_para["type"]

        # 選擇模板
        if para_type == "heading":
            level = md_para.get("level", 1)
            if level == 1:
                template = heading1_template
            elif level == 2:
                template = heading2_template
            else:
                template = heading3_template
        else:
            template = normal_template

        # 克隆並追加到末尾
        new_elem = deepcopy(template._element)
        body.append(new_elem)

        # 設定文字
        t_elements = list(new_elem.iter(qn("w:t")))
        if t_elements:
            for t in t_elements:
                t.text = ""
            t_elements[0].text = text

        insert_count += 1

    # 儲存
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))

    print(f"\n=== 完成 ===")
    print(f"  - 插入段落：{insert_count}")
    print(f"  - 已儲存：{args.output}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
