#!/usr/bin/env python3
"""
sync_docx_content.py - 保留格式的 DOCX 內容同步

安全策略：
1. 只沿既有段落順序覆寫文字內容，完全保留格式
2. 不新增、不刪除任何段落
3. 輸出需要手動處理的差異報告
4. 僅適用於已人工確認章節與段落順序一致的小幅同步

用法：
    python sync_docx_content.py template.docx content.md output.docx
    python sync_docx_content.py template.docx content.md output.docx --report diff.txt
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from markdown_it import MarkdownIt


def parse_args():
    parser = argparse.ArgumentParser(description="同步 MD 內容到 DOCX，保留原始格式；僅適用於結構已人工確認的一致版本")
    parser.add_argument("template", type=Path, help="原始 DOCX 模板")
    parser.add_argument("markdown", type=Path, help="Markdown 內容來源")
    parser.add_argument("output", type=Path, help="輸出 DOCX")
    parser.add_argument("--start-heading", default="中文摘要", help="從此標題開始同步")
    parser.add_argument("--report", type=Path, help="輸出差異報告")
    parser.add_argument("--dry-run", action="store_true", help="只顯示變更，不實際修改")
    return parser.parse_args()


def get_paragraph_text(para) -> str:
    """取得段落的純文字內容"""
    return "".join(node.text or "" for node in para._element.iter(qn("w:t"))).strip()


def set_paragraph_text(para, new_text: str) -> bool:
    """
    替換段落文字，完全保留格式。
    清空所有 <w:t>，在第一個 <w:t> 填入新文字。
    """
    t_elements = list(para._element.iter(qn("w:t")))
    if not t_elements:
        return False

    for t in t_elements:
        t.text = ""
    t_elements[0].text = new_text
    return True


def normalize_text(text: str) -> str:
    """正規化文字以便比較"""
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[「」『』【】（）\(\)]", "", text)
    return text.strip().lower()


def normalize_heading(text: str) -> str:
    """正規化標題以便匹配"""
    text = re.sub(r"^[\d\.]+\s*", "", text)
    text = re.sub(r"^第[一二三四五六七八九十]+[章節]\s*", "", text)
    text = re.sub(r"^\#{1,6}\s*", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def is_heading_paragraph(para) -> bool:
    """判斷是否為標題段落"""
    style = para.style
    if style and style.name and "Heading" in style.name:
        return True
    text = get_paragraph_text(para)
    if re.match(r"^(第[一二三四五六七八九十]+[章節]|[\d\.]+\s+\S)", text):
        return True
    return False


def inline_text(token) -> str:
    """從 token 提取純文字"""
    parts = []
    for child in token.children or []:
        if child.type in {"text", "code_inline"}:
            parts.append(child.content)
        elif child.type in {"softbreak", "hardbreak"}:
            parts.append(" ")
    return "".join(parts).strip()


def parse_markdown(md_text: str, start_heading: str) -> list[dict]:
    """解析 Markdown 為段落列表"""
    md = MarkdownIt("commonmark", {"html": True}).enable("table")
    tokens = md.parse(md_text)

    blocks = []
    started = False
    i = 0

    while i < len(tokens):
        token = tokens[i]

        if token.type == "heading_open":
            level = int(token.tag[1])
            text = inline_text(tokens[i + 1])
            if not started and text == start_heading:
                started = True
            if started:
                blocks.append({"type": "heading", "level": level, "text": text})
            i += 3
            continue

        if not started:
            i += 1
            continue

        if token.type == "paragraph_open":
            text = inline_text(tokens[i + 1])
            if text:
                blocks.append({"type": "paragraph", "text": text})
            i += 3
            continue

        if token.type in {"ordered_list_open", "bullet_list_open"}:
            ordered = token.type == "ordered_list_open"
            close_type = "ordered_list_close" if ordered else "bullet_list_close"
            i += 1
            item_index = 1
            while i < len(tokens) and tokens[i].type != close_type:
                if tokens[i].type == "list_item_open":
                    i += 1
                    item_parts = []
                    while i < len(tokens) and tokens[i].type != "list_item_close":
                        if tokens[i].type == "paragraph_open":
                            text = inline_text(tokens[i + 1])
                            if text:
                                item_parts.append(text)
                            i += 3
                            continue
                        i += 1
                    item_text = " ".join(item_parts).strip()
                    prefix = f"{item_index}. " if ordered else "• "
                    blocks.append({"type": "list_item", "text": prefix + item_text})
                    item_index += 1
                i += 1
            i += 1
            continue

        if token.type == "fence":
            blocks.append({"type": "code", "text": token.content.strip()})
            i += 1
            continue

        i += 1

    return blocks


def find_start_index(paragraphs: list, start_heading: str) -> int:
    """找到起始標題的段落索引"""
    for i, para in enumerate(paragraphs):
        if get_paragraph_text(para) == start_heading:
            return i
    return -1


def build_heading_map(paragraphs: list, start_idx: int) -> dict:
    """建立標題映射：正規化標題 -> [(原始文字, 段落索引)]"""
    mapping = {}
    for i, para in enumerate(paragraphs[start_idx:], start=start_idx):
        text = get_paragraph_text(para)
        if text and is_heading_paragraph(para):
            norm = normalize_heading(text)
            if norm not in mapping:
                mapping[norm] = []
            mapping[norm].append((text, i))
    return mapping


def sync_content(doc: Document, md_blocks: list[dict], start_idx: int, dry_run: bool = False):
    """
    同步內容，只替換文字，不增刪段落
    回傳：(changes, manual_items)
    """
    paragraphs = doc.paragraphs
    heading_map = build_heading_map(paragraphs, start_idx)

    changes = []
    manual_items = []  # 需要手動處理的項目

    docx_idx = start_idx
    heading_used = set()  # 已使用的標題索引

    for md_block in md_blocks:
        md_text = md_block["text"]
        md_type = md_block["type"]

        if md_type == "heading":
            norm = normalize_heading(md_text)

            # 查找匹配的標題
            matched = False
            if norm in heading_map:
                for orig_text, para_idx in heading_map[norm]:
                    if para_idx not in heading_used and para_idx >= docx_idx:
                        heading_used.add(para_idx)
                        docx_idx = para_idx

                        # 如果標題文字不同，更新
                        if orig_text != md_text:
                            if not dry_run:
                                set_paragraph_text(paragraphs[para_idx], md_text)
                            changes.append(("UPDATE_HEADING", para_idx, f"'{orig_text[:35]}' -> '{md_text[:35]}'"))
                        else:
                            changes.append(("MATCH_HEADING", para_idx, md_text[:50]))

                        docx_idx += 1
                        matched = True
                        break

            if not matched:
                manual_items.append(("MISSING_HEADING", md_text))

        else:
            # 非標題段落：找下一個非標題段落替換
            found = False
            search_limit = min(docx_idx + 10, len(paragraphs))

            for search_idx in range(docx_idx, search_limit):
                para = paragraphs[search_idx]
                if is_heading_paragraph(para):
                    continue

                docx_text = get_paragraph_text(para)
                if docx_text:  # 有內容的段落
                    if not dry_run:
                        set_paragraph_text(para, md_text)
                    changes.append(("REPLACE", search_idx, f"'{docx_text[:20]}...' -> '{md_text[:20]}...'"))
                    docx_idx = search_idx + 1
                    found = True
                    break

            if not found:
                manual_items.append(("MISSING_PARA", md_text[:80]))

    return changes, manual_items


def main():
    args = parse_args()

    if not args.template.exists():
        print(f"錯誤：找不到模板 {args.template}", file=sys.stderr)
        return 1

    if not args.markdown.exists():
        print(f"錯誤：找不到 Markdown {args.markdown}", file=sys.stderr)
        return 1

    # 解析 Markdown
    md_text = args.markdown.read_text(encoding="utf-8")
    md_blocks = parse_markdown(md_text, args.start_heading)
    print(f"解析 Markdown：{len(md_blocks)} 個區塊")

    # 載入 DOCX
    doc = Document(str(args.template))
    start_idx = find_start_index(doc.paragraphs, args.start_heading)

    if start_idx == -1:
        print(f"錯誤：在 DOCX 中找不到起始標題 '{args.start_heading}'", file=sys.stderr)
        return 1

    print(f"起始位置：段落 {start_idx}")
    print(f"DOCX 總段落數：{len(doc.paragraphs)}")

    # 同步內容
    changes, manual_items = sync_content(doc, md_blocks, start_idx, args.dry_run)

    # 統計
    replace_count = sum(1 for c in changes if c[0] == "REPLACE")
    heading_count = sum(1 for c in changes if "HEADING" in c[0])

    print(f"\n變更統計：")
    print(f"  - 標題匹配/更新：{heading_count}")
    print(f"  - 段落替換：{replace_count}")
    print(f"  - 需手動處理：{len(manual_items)}")

    # 顯示前 30 個變更
    print(f"\n變更詳情（前 30 項）：")
    for change_type, idx, desc in changes[:30]:
        print(f"  [{change_type}] #{idx}: {desc}")
    if len(changes) > 30:
        print(f"  ... 還有 {len(changes) - 30} 項")

    # 需手動處理的項目
    if manual_items:
        print(f"\n需手動處理（{len(manual_items)} 項）：")
        for item_type, text in manual_items[:20]:
            print(f"  [{item_type}] {text[:60]}...")
        if len(manual_items) > 20:
            print(f"  ... 還有 {len(manual_items) - 20} 項")

    # 輸出報告
    if args.report:
        with open(args.report, "w", encoding="utf-8") as f:
            f.write("# DOCX 同步差異報告\n\n")
            f.write(f"來源：{args.template} + {args.markdown}\n")
            f.write(f"輸出：{args.output}\n\n")

            f.write("## 已處理變更\n\n")
            for change_type, idx, desc in changes:
                f.write(f"- [{change_type}] #{idx}: {desc}\n")

            f.write("\n## 需手動處理\n\n")
            for item_type, text in manual_items:
                f.write(f"- [{item_type}] {text}\n")

        print(f"\n報告已儲存：{args.report}")

    # 儲存
    if args.dry_run:
        print("\n(dry-run 模式，未實際修改)")
    else:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(args.output))
        print(f"\n已儲存：{args.output}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
