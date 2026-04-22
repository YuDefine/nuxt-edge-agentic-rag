#!/usr/bin/env python3
"""
clone_section.py - 複製 DOCX 章節結構並插入新內容

用法：
    python clone_section.py input.docx output.docx --source "第二節" --target "第三節" --content "新內容..."
    python clone_section.py input.docx output.docx --source "### 2.2.1" --target "### 2.2.3" --content-file new_section.md

此工具會：
1. 找到 source 標題所在的段落
2. 複製該段落的完整格式（字體、樣式、編號）
3. 將文字替換為 target 標題
4. 在正確位置插入新段落
5. 若提供 content，則在標題後插入內容段落（複製相鄰內容段落的格式）
"""

import argparse
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def find_paragraph_by_text(doc: Document, search_text: str) -> tuple[int, any]:
    """找到包含指定文字的段落，回傳 (index, paragraph)"""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            return i, para
    return -1, None


def find_section_range(doc: Document, start_idx: int, heading_style: str) -> tuple[int, int]:
    """
    找到章節的範圍（從標題到下一個同級或更高級標題之前）
    回傳 (start_idx, end_idx)，end_idx 是下一個章節的起始位置
    """
    end_idx = len(doc.paragraphs)
    start_style = doc.paragraphs[start_idx].style.name if doc.paragraphs[start_idx].style else None

    for i in range(start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else None

        # 檢查是否為同級或更高級標題
        if style_name and 'Heading' in style_name:
            # 比較標題層級
            if style_name <= start_style if start_style else True:
                end_idx = i
                break

        # 也檢查文字模式（第X節、第X章）
        text = para.text.strip()
        if text.startswith('第') and ('章' in text or '節' in text):
            end_idx = i
            break
        if text.startswith('##'):
            end_idx = i
            break

    return start_idx, end_idx


def clone_paragraph(source_para) -> any:
    """深複製段落的 XML 結構"""
    new_elem = deepcopy(source_para._element)
    return new_elem


def insert_paragraph_after(doc: Document, ref_para, new_elem):
    """在參考段落後插入新段落元素"""
    ref_para._element.addnext(new_elem)


def replace_paragraph_text(para_elem, new_text: str):
    """替換段落元素中的文字，保留格式"""
    # 找到所有 run 中的 text 元素
    for t in para_elem.iter(qn('w:t')):
        t.text = ''

    # 在第一個 run 中設定新文字
    runs = para_elem.findall('.//' + qn('w:r'))
    if runs:
        t_elem = runs[0].find(qn('w:t'))
        if t_elem is not None:
            t_elem.text = new_text
        else:
            # 建立新的 text 元素
            t_elem = OxmlElement('w:t')
            t_elem.text = new_text
            runs[0].append(t_elem)


def clone_section(
    input_path: str,
    output_path: str,
    source_text: str,
    target_text: str,
    content: str = None,
    content_file: str = None,
    insert_after: str = None
):
    """
    複製章節結構並插入新內容

    Args:
        input_path: 來源 DOCX
        output_path: 輸出 DOCX
        source_text: 要複製的來源章節標題文字
        target_text: 新章節的標題文字
        content: 新章節的內容（純文字或 Markdown）
        content_file: 從檔案讀取內容
        insert_after: 插入位置（在此標題之後），若不指定則在 source 之後
    """
    doc = Document(input_path)

    # 找到來源段落
    source_idx, source_para = find_paragraph_by_text(doc, source_text)
    if source_idx == -1:
        print(f"錯誤：找不到包含 '{source_text}' 的段落", file=sys.stderr)
        sys.exit(1)

    print(f"找到來源段落：index={source_idx}, text='{source_para.text[:50]}...'")

    # 決定插入位置
    if insert_after:
        insert_idx, insert_para = find_paragraph_by_text(doc, insert_after)
        if insert_idx == -1:
            print(f"錯誤：找不到插入位置 '{insert_after}'", file=sys.stderr)
            sys.exit(1)
    else:
        # 找到來源章節的結束位置
        _, section_end = find_section_range(doc, source_idx, source_para.style.name if source_para.style else '')
        insert_idx = section_end - 1
        insert_para = doc.paragraphs[insert_idx]

    print(f"插入位置：index={insert_idx}, after='{insert_para.text[:50]}...'")

    # 複製標題段落
    new_heading_elem = clone_paragraph(source_para)
    replace_paragraph_text(new_heading_elem, target_text)

    # 讀取內容
    if content_file:
        content = Path(content_file).read_text(encoding='utf-8')

    # 找到來源章節的第一個內容段落（用於複製格式）
    content_template_para = None
    if source_idx + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[source_idx + 1]
        if next_para.style and 'Heading' not in next_para.style.name:
            content_template_para = next_para

    # 插入標題
    insert_paragraph_after(doc, insert_para, new_heading_elem)

    # 插入內容（如果有）
    if content and content_template_para:
        # 分割內容為段落
        content_paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]

        # 從後往前插入，這樣順序才會正確
        last_inserted = new_heading_elem
        for para_text in content_paragraphs:
            new_content_elem = clone_paragraph(content_template_para)
            replace_paragraph_text(new_content_elem, para_text)
            last_inserted.addnext(new_content_elem)
            last_inserted = new_content_elem

    # 儲存
    doc.save(output_path)
    print(f"已儲存：{output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='複製 DOCX 章節結構並插入新內容',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
範例：
  # 複製「第二節」的格式，建立「第三節」
  python clone_section.py report.docx report_new.docx --source "第二節" --target "第三節"

  # 複製並加入內容
  python clone_section.py report.docx report_new.docx \\
    --source "### 2.2.1" --target "### 2.2.3 新小節" \\
    --content "這是新小節的內容。"

  # 從檔案讀取內容
  python clone_section.py report.docx report_new.docx \\
    --source "第一節" --target "第四節" \\
    --content-file new_section.txt
"""
    )

    parser.add_argument('input', help='來源 DOCX 檔案')
    parser.add_argument('output', help='輸出 DOCX 檔案')
    parser.add_argument('--source', required=True, help='要複製的來源章節標題（部分匹配）')
    parser.add_argument('--target', required=True, help='新章節的標題')
    parser.add_argument('--content', help='新章節的內容')
    parser.add_argument('--content-file', help='從檔案讀取內容')
    parser.add_argument('--insert-after', help='插入位置（在此標題之後）')

    args = parser.parse_args()

    clone_section(
        args.input,
        args.output,
        args.source,
        args.target,
        args.content,
        args.content_file,
        args.insert_after
    )


if __name__ == '__main__':
    main()
