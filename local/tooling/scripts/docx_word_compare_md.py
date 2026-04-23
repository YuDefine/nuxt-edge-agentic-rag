#!/usr/bin/env python3
"""
docx_word_compare_md.py - 產生 Word 手動回填用的對照 Markdown

用途：
1. 以 Markdown 為目標內容來源，方便逐段複製貼上到 Word。
2. 同時附上目前 DOCX 的純文字段落，方便人工對照。
3. 不直接修改 DOCX；只生成一份對照稿。

用法：
    python local/tooling/scripts/docx_word_compare_md.py \
      local/reports/archive/latest.docx \
      local/reports/latest.md \
      local/reports/notes/latest-word-compare.md
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
DOCX_NUMERIC_HEADING_RE = re.compile(r"^(?:\d+(?:\.\d+){0,4})\s+")
DOCX_CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百零〇兩]+章\s+")
DOCX_SECTION_RE = re.compile(r"^第[一二三四五六七八九十百零〇兩]+節\s+")
DOCX_APPENDIX_RE = re.compile(r"^附錄\s*[A-ZＡ-Ｚ][：: ]*")
WS_RE = re.compile(r"\s+")


@dataclass
class MarkdownBlock:
    index: int
    title: str
    heading_line: str
    level: int
    body_lines: list[str]

    @property
    def copy_text(self) -> str:
        lines = [self.heading_line]
        if self.body_lines:
            lines.append("\n".join(self.body_lines).rstrip())
        return "\n\n".join(line for line in lines if line).strip()


@dataclass
class DocxSection:
    title: str
    paragraphs: list[str]


def get_docx_paragraph_text(paragraph) -> str:
    return "".join(node.text or "" for node in paragraph._element.iter(qn("w:t"))).strip()


def get_docx_heading_title(paragraph) -> str | None:
    style_name = paragraph.style.name if paragraph.style else ""
    text = get_docx_paragraph_text(paragraph)
    if not text:
        return None

    if style_name and "Heading" in style_name:
        return text

    if DOCX_NUMERIC_HEADING_RE.match(text):
        return text

    if DOCX_CHAPTER_RE.match(text):
        return text

    if DOCX_SECTION_RE.match(text):
        return text

    if DOCX_APPENDIX_RE.match(text):
        return text

    if text in {"中文摘要", "目錄", "符號與用詞索引", "圖表目錄", "參考文獻", "附錄"}:
        return text

    return None


def normalize_heading(value: str) -> str:
    value = value.strip()
    value = re.sub(r"^#{1,6}\s+", "", value)
    value = DOCX_CHAPTER_RE.sub("", value)
    value = DOCX_SECTION_RE.sub("", value)
    value = DOCX_APPENDIX_RE.sub("", value)
    value = DOCX_NUMERIC_HEADING_RE.sub("", value)
    value = WS_RE.sub(" ", value)
    return value.strip().lower()


def slugify(value: str) -> str:
    value = normalize_heading(value)
    value = re.sub(r"[^\w\u4e00-\u9fff\- ]+", "", value)
    value = value.replace(" ", "-")
    return value or "section"


def parse_markdown_blocks(markdown_path: Path) -> list[MarkdownBlock]:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()

    blocks: list[MarkdownBlock] = []
    preamble: list[str] = []
    current_heading_line: str | None = None
    current_title: str | None = None
    current_level: int | None = None
    current_body: list[str] = []
    block_index = 1

    for line in lines:
        match = HEADING_RE.match(line)
        if match:
            if current_heading_line is None:
                if preamble:
                    blocks.append(
                        MarkdownBlock(
                            index=block_index,
                            title="封面與前置內容",
                            heading_line="<!-- 前置內容 -->",
                            level=0,
                            body_lines=preamble[:],
                        )
                    )
                    block_index += 1
                    preamble.clear()
            else:
                blocks.append(
                    MarkdownBlock(
                        index=block_index,
                        title=current_title or "未命名段落",
                        heading_line=current_heading_line,
                        level=current_level or 1,
                        body_lines=current_body[:],
                    )
                )
                block_index += 1

            current_heading_line = line
            current_title = match.group(2).strip()
            current_level = len(match.group(1))
            current_body = []
            continue

        if current_heading_line is None:
            preamble.append(line)
        else:
            current_body.append(line)

    if current_heading_line is None:
        blocks.append(
            MarkdownBlock(
                index=block_index,
                title="全文",
                heading_line="<!-- 全文 -->",
                level=0,
                body_lines=lines,
            )
        )
    else:
        blocks.append(
            MarkdownBlock(
                index=block_index,
                title=current_title or "未命名段落",
                heading_line=current_heading_line,
                level=current_level or 1,
                body_lines=current_body[:],
            )
        )

    return blocks


def parse_docx_sections(docx_path: Path) -> dict[str, list[DocxSection]]:
    document = Document(str(docx_path))
    sections: list[DocxSection] = []

    preamble: list[str] = []
    current_title: str | None = None
    current_paragraphs: list[str] = []

    for paragraph in document.paragraphs:
        text = get_docx_paragraph_text(paragraph)
        if not text:
            continue

        heading_title = get_docx_heading_title(paragraph)
        if heading_title:
            if current_title is None:
                if preamble:
                    sections.append(DocxSection(title="封面與前置內容", paragraphs=preamble[:]))
                    preamble.clear()
            else:
                sections.append(DocxSection(title=current_title, paragraphs=current_paragraphs[:]))

            current_title = heading_title
            current_paragraphs = [heading_title]
            continue

        if current_title is None:
            preamble.append(text)
        else:
            current_paragraphs.append(text)

    if current_title is None:
        if preamble:
            sections.append(DocxSection(title="封面與前置內容", paragraphs=preamble[:]))
    else:
        sections.append(DocxSection(title=current_title, paragraphs=current_paragraphs[:]))

    grouped: dict[str, list[DocxSection]] = {}
    for section in sections:
        key = normalize_heading(section.title)
        grouped.setdefault(key, []).append(section)
    return grouped


def render_docx_section(section: DocxSection | None) -> str:
    if section is None:
        return "找不到同名 DOCX 章節，請人工比對附近位置。"
    return "\n\n".join(section.paragraphs).strip()


def iter_index_lines(blocks: Iterable[MarkdownBlock]) -> list[str]:
    lines = []
    for block in blocks:
        anchor = slugify(f"{block.index:04d}-{block.title}")
        lines.append(f"- [{block.index:04d} {block.title}](#{anchor})")
    return lines


def build_compare_markdown(
    *,
    docx_path: Path,
    markdown_path: Path,
    blocks: list[MarkdownBlock],
    docx_sections: dict[str, list[DocxSection]],
    copy_only: bool,
) -> str:
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if copy_only:
        lines = [
            "# Word 純複製貼上稿",
            "",
            f"- 來源 MD：`{markdown_path}`",
            f"- 生成時間：`{generated_at}`",
            "",
            "## 使用方式",
            "",
            "1. 直接複製各段的 Markdown 內容。",
            "2. 貼到 Word 後只處理排版，不在這份稿內做比對。",
            "3. 表格、圖片、目錄、頁碼、圖說仍以 Word 手動調整為準。",
            "",
            "## 段落索引",
            "",
            *iter_index_lines(blocks),
            "",
        ]
    else:
        lines = [
            "# Word 手動回填對照稿",
            "",
            f"- 來源 DOCX：`{docx_path}`",
            f"- 來源 MD：`{markdown_path}`",
            f"- 生成時間：`{generated_at}`",
            "",
            "## 使用方式",
            "",
            "1. 先看每段的「目標內容（可直接複製）」。",
            "2. 貼到 Word 前，先對照同段的「目前 DOCX 內容」。",
            "3. 表格、圖片、目錄、頁碼、圖說仍以 Word 手動調整為準。",
            "",
            "## 段落索引",
            "",
            *iter_index_lines(blocks),
            "",
        ]

    for block in blocks:
        key = normalize_heading(block.title)
        matched_sections = docx_sections.get(key, [])
        matched_section = matched_sections[0] if matched_sections else None
        status = "DOCX 有同名段落" if matched_section else "DOCX 無同名段落，需人工定位"
        anchor_title = f"{block.index:04d} {block.title}"

        if copy_only:
            lines.extend(
                [
                    "---",
                    "",
                    f"## {anchor_title}",
                    "",
                    "```markdown",
                    block.copy_text,
                    "```",
                    "",
                ]
            )
        else:
            lines.extend(
                [
                    "---",
                    "",
                    f"## {anchor_title}",
                    "",
                    f"- 狀態：{status}",
                    f"- 原始標題：`{block.heading_line}`",
                    "",
                    "### 目標內容（可直接複製）",
                    "",
                    "```markdown",
                    block.copy_text,
                    "```",
                    "",
                    "### 目前 DOCX 內容（對照）",
                    "",
                    "```text",
                    render_docx_section(matched_section),
                    "```",
                    "",
                ]
            )

    return "\n".join(lines).rstrip() + "\n"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="產生 Word 手動回填對照 Markdown")
    parser.add_argument("docx", type=Path, help="來源 DOCX")
    parser.add_argument("markdown", type=Path, help="來源 Markdown")
    parser.add_argument("output", type=Path, help="輸出 Markdown")
    parser.add_argument("--copy-only", action="store_true", help="只輸出可複製內容，不附 DOCX 對照")
    return parser.parse_args()


def main() -> int:
    args = parse_args()

    if not args.docx.exists():
        raise SystemExit(f"找不到 DOCX：{args.docx}")
    if not args.markdown.exists():
        raise SystemExit(f"找不到 Markdown：{args.markdown}")

    blocks = parse_markdown_blocks(args.markdown)
    docx_sections = parse_docx_sections(args.docx)
    output = build_compare_markdown(
        docx_path=args.docx,
        markdown_path=args.markdown,
        blocks=blocks,
        docx_sections=docx_sections,
        copy_only=args.copy_only,
    )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(output, encoding="utf-8")

    print(f"已產生：{args.output}")
    print(f"Markdown 區塊數：{len(blocks)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
