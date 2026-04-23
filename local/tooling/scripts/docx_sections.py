#!/usr/bin/env python3
"""
docx_sections.py - 按章節分析和處理 DOCX

用法：
    # 列出 DOCX 所有章節
    python docx_sections.py list template.docx

    # 列出特定章節的段落
    python docx_sections.py show template.docx "1.1.1"

    # 比較 DOCX 和 MD 的特定章節
    python docx_sections.py compare template.docx content.md "1.1.1"

    # 更新特定章節的內容
    python docx_sections.py update template.docx content.md output.docx "1.1.1"
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
from markdown_it import MarkdownIt


def get_para_text(para) -> str:
    return "".join(node.text or "" for node in para._element.iter(qn("w:t"))).strip()


def set_para_text(para, new_text: str) -> bool:
    t_elements = list(para._element.iter(qn("w:t")))
    if not t_elements:
        return False
    for t in t_elements:
        t.text = ""
    t_elements[0].text = new_text
    return True


def clone_para_after(para, new_text: str):
    new_elem = deepcopy(para._element)
    para._element.addnext(new_elem)
    t_elements = list(new_elem.iter(qn("w:t")))
    if t_elements:
        for t in t_elements:
            t.text = ""
        t_elements[0].text = new_text
    return new_elem


def get_heading_id(text: str) -> str | None:
    """提取章節編號，如 '1.1.1'、'第一節'"""
    # 數字編號
    match = re.match(r"^([\d\.]+)", text)
    if match:
        return match.group(1).rstrip(".")

    # 中文編號
    match = re.match(r"^(第[一二三四五六七八九十]+[章節])", text)
    if match:
        return match.group(1)

    return None


def get_heading_level(para) -> int:
    style = para.style
    if style and style.name:
        match = re.search(r"Heading\s*(\d+)", style.name)
        if match:
            return int(match.group(1))
    return 0


def is_heading(para) -> bool:
    level = get_heading_level(para)
    if level > 0:
        return True
    text = get_para_text(para)
    return get_heading_id(text) is not None


def find_section_range(paragraphs: list, section_id: str) -> tuple[int, int] | None:
    """找到章節的起始和結束索引 [start, end)"""
    start_idx = None
    start_level = 0

    for i, para in enumerate(paragraphs):
        text = get_para_text(para)
        para_id = get_heading_id(text)

        if para_id and section_id in para_id:
            start_idx = i
            start_level = get_heading_level(para)
            break

    if start_idx is None:
        return None

    # 找結束位置（下一個同級或更高級的標題）
    end_idx = len(paragraphs)
    for i in range(start_idx + 1, len(paragraphs)):
        para = paragraphs[i]
        level = get_heading_level(para)
        if level > 0 and level <= start_level:
            end_idx = i
            break

    return (start_idx, end_idx)


def list_sections(docx_path: str):
    """列出所有章節"""
    doc = Document(docx_path)
    sections = []

    for i, para in enumerate(doc.paragraphs):
        text = get_para_text(para)
        if not text:
            continue

        level = get_heading_level(para)
        heading_id = get_heading_id(text)

        if level > 0 or heading_id:
            indent = "  " * (level - 1) if level > 0 else ""
            sections.append({
                "index": i,
                "level": level,
                "id": heading_id,
                "text": text[:60]
            })
            print(f"{indent}[{i}] {text[:60]}{'...' if len(text) > 60 else ''}")

    return sections


def show_section(docx_path: str, section_id: str):
    """顯示特定章節的所有段落"""
    doc = Document(docx_path)
    range_result = find_section_range(doc.paragraphs, section_id)

    if not range_result:
        print(f"找不到章節：{section_id}")
        return

    start, end = range_result
    print(f"章節範圍：段落 {start} - {end} ({end - start} 段)")
    print("=" * 60)

    for i in range(start, end):
        para = doc.paragraphs[i]
        text = get_para_text(para)
        style = para.style.name if para.style else "Normal"
        level = get_heading_level(para)

        prefix = f"[H{level}]" if level > 0 else "[P]"
        print(f"{i:3d} {prefix:5s} {style:20s} | {text[:50]}{'...' if len(text) > 50 else ''}")


def inline_text(token) -> str:
    parts = []
    for child in token.children or []:
        if child.type in {"text", "code_inline"}:
            parts.append(child.content)
        elif child.type in {"softbreak", "hardbreak"}:
            parts.append(" ")
    return "".join(parts).strip()


def parse_md_section(md_path: str, section_id: str) -> list[dict] | None:
    """從 MD 提取特定章節的段落"""
    md_text = Path(md_path).read_text(encoding="utf-8")
    md = MarkdownIt("commonmark", {"html": True}).enable("table")
    tokens = md.parse(md_text)

    # 找到章節起始
    in_section = False
    section_level = 0
    blocks = []
    i = 0

    while i < len(tokens):
        token = tokens[i]

        if token.type == "heading_open":
            level = int(token.tag[1])
            text = inline_text(tokens[i + 1])
            heading_id = get_heading_id(text)

            if heading_id and section_id in heading_id:
                in_section = True
                section_level = level
                blocks.append({"type": "heading", "level": level, "text": text})
                i += 3
                continue

            if in_section and level <= section_level:
                # 遇到同級或更高級標題，結束
                break

            if in_section:
                blocks.append({"type": "heading", "level": level, "text": text})

            i += 3
            continue

        if not in_section:
            i += 1
            continue

        if token.type == "paragraph_open":
            text = inline_text(tokens[i + 1])
            if text:
                blocks.append({"type": "paragraph", "text": text})
            i += 3
            continue

        if token.type in {"ordered_list_open", "bullet_list_open"}:
            ordered = token.type == "ordered_list_open"
            close_type = "ordered_list_close" if ordered else "bullet_list_close"
            i += 1
            item_index = 1
            while i < len(tokens) and tokens[i].type != close_type:
                if tokens[i].type == "list_item_open":
                    i += 1
                    item_parts = []
                    while i < len(tokens) and tokens[i].type != "list_item_close":
                        if tokens[i].type == "paragraph_open":
                            text = inline_text(tokens[i + 1])
                            if text:
                                item_parts.append(text)
                            i += 3
                            continue
                        i += 1
                    item_text = " ".join(item_parts).strip()
                    prefix = f"{item_index}. " if ordered else "• "
                    blocks.append({"type": "list_item", "text": prefix + item_text})
                    item_index += 1
                i += 1
            i += 1
            continue

        i += 1

    return blocks if blocks else None


def compare_section(docx_path: str, md_path: str, section_id: str):
    """比較特定章節"""
    doc = Document(docx_path)
    docx_range = find_section_range(doc.paragraphs, section_id)

    if not docx_range:
        print(f"DOCX 中找不到章節：{section_id}")
        return

    md_blocks = parse_md_section(md_path, section_id)
    if not md_blocks:
        print(f"MD 中找不到章節：{section_id}")
        return

    start, end = docx_range
    docx_paras = []
    for i in range(start, end):
        text = get_para_text(doc.paragraphs[i])
        if text:
            docx_paras.append({"index": i, "text": text})

    print(f"DOCX 段落數：{len(docx_paras)}")
    print(f"MD 段落數：{len(md_blocks)}")
    print()

    # 並排顯示
    max_len = max(len(docx_paras), len(md_blocks))
    print(f"{'DOCX':<40} | {'MD':<40}")
    print("-" * 83)

    for i in range(max_len):
        docx_text = docx_paras[i]["text"][:38] if i < len(docx_paras) else ""
        md_text = md_blocks[i]["text"][:38] if i < len(md_blocks) else ""
        print(f"{docx_text:<40} | {md_text:<40}")


def update_section(docx_path: str, md_path: str, output_path: str, section_id: str, dry_run: bool = False):
    """更新特定章節的內容"""
    doc = Document(docx_path)
    docx_range = find_section_range(doc.paragraphs, section_id)

    if not docx_range:
        print(f"DOCX 中找不到章節：{section_id}")
        return False

    md_blocks = parse_md_section(md_path, section_id)
    if not md_blocks:
        print(f"MD 中找不到章節：{section_id}")
        return False

    start, end = docx_range
    docx_paras = doc.paragraphs[start:end]

    # 過濾掉空段落
    docx_with_content = [(i + start, p) for i, p in enumerate(docx_paras) if get_para_text(p)]

    print(f"DOCX 章節段落：{len(docx_with_content)}")
    print(f"MD 章節段落：{len(md_blocks)}")

    # 逐一替換
    changes = []
    for i, (docx_idx, para) in enumerate(docx_with_content):
        if i < len(md_blocks):
            old_text = get_para_text(para)
            new_text = md_blocks[i]["text"]

            if old_text != new_text:
                changes.append({
                    "action": "REPLACE",
                    "index": docx_idx,
                    "old": old_text[:50],
                    "new": new_text[:50]
                })
                if not dry_run:
                    set_para_text(para, new_text)

    # 需要插入的段落
    if len(md_blocks) > len(docx_with_content):
        for i in range(len(docx_with_content), len(md_blocks)):
            changes.append({
                "action": "INSERT",
                "text": md_blocks[i]["text"][:50]
            })
            if not dry_run and docx_with_content:
                # 在最後一個段落後插入
                last_idx, last_para = docx_with_content[-1]
                clone_para_after(last_para, md_blocks[i]["text"])

    print(f"\n變更：{len(changes)}")
    for c in changes[:20]:
        if c["action"] == "REPLACE":
            print(f"  [REPLACE] #{c['index']}: '{c['old']}...' -> '{c['new']}...'")
        else:
            print(f"  [INSERT]: {c['text']}...")

    if len(changes) > 20:
        print(f"  ... 還有 {len(changes) - 20} 項")

    if dry_run:
        print("\n(dry-run 模式)")
    else:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        print(f"\n已儲存：{output_path}")

    return True


def main():
    parser = argparse.ArgumentParser(description="按章節處理 DOCX")
    subparsers = parser.add_subparsers(dest="command", required=True)

    # list
    list_parser = subparsers.add_parser("list", help="列出所有章節")
    list_parser.add_argument("docx", type=Path)

    # show
    show_parser = subparsers.add_parser("show", help="顯示特定章節")
    show_parser.add_argument("docx", type=Path)
    show_parser.add_argument("section", help="章節編號，如 1.1.1")

    # compare
    compare_parser = subparsers.add_parser("compare", help="比較章節")
    compare_parser.add_argument("docx", type=Path)
    compare_parser.add_argument("markdown", type=Path)
    compare_parser.add_argument("section", help="章節編號")

    # update
    update_parser = subparsers.add_parser("update", help="更新章節")
    update_parser.add_argument("docx", type=Path)
    update_parser.add_argument("markdown", type=Path)
    update_parser.add_argument("output", type=Path)
    update_parser.add_argument("section", help="章節編號")
    update_parser.add_argument("--dry-run", action="store_true")

    args = parser.parse_args()

    if args.command == "list":
        list_sections(str(args.docx))
    elif args.command == "show":
        show_section(str(args.docx), args.section)
    elif args.command == "compare":
        compare_section(str(args.docx), str(args.markdown), args.section)
    elif args.command == "update":
        update_section(str(args.docx), str(args.markdown), str(args.output), args.section, args.dry_run)


if __name__ == "__main__":
    main()
